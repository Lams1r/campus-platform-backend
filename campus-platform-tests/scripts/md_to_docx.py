from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TEST_SUMMARY_REPORT.md"
OUTPUT = ROOT / "TEST_SUMMARY_REPORT.docx"


def set_run_font(run, *, size: int = 11, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 6) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_inline_runs(paragraph, text: str, *, size: int = 11) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        set_run_font(run, size=size)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.color.rgb = RGBColor(59, 76, 184)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, value.strip(), size=10)
            for run in paragraph.runs:
                run.bold = row_index == 0
            if row_index == 0:
                shade_cell(cell, "EAF0FF")
    document.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code_block = False
    code_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code_block:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.4)
                set_paragraph_spacing(paragraph, after=8)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(9)
                code_lines.clear()
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        if line.strip().startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, before=4, after=18)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=20, bold=True, color="172554")
            index += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            paragraph = document.add_heading(level=level)
            paragraph.clear()
            set_paragraph_spacing(paragraph, before=10, after=6)
            run = paragraph.add_run(heading_match.group(2).strip())
            set_run_font(run, size=max(12, 17 - level), bold=True, color="1E3A8A")
            index += 1
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            set_paragraph_spacing(paragraph, after=3)
            add_inline_runs(paragraph, line[2:].strip())
            index += 1
            continue

        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph)
        add_inline_runs(paragraph, line.strip())
        index += 1

    document.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
